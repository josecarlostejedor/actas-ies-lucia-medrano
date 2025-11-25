import streamlit as st
import subprocess
import sys
import time
import os

# --- BLOQUE DE AUTO-INSTALACIÓN ---
# Garantiza que las dependencias estén presentes en el entorno volátil de la nube
def instalar(package):
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
    except:
        pass 

try:
    import openai
    from docx import Document
    from fpdf import FPDF
    # Verificación estricta de versión para compatibilidad con st.audio_input
    ver = st.__version__.split('.')
    if int(ver[0]) < 1 or (int(ver[0]) == 1 and int(ver[1]) < 40):
        raise ImportError("Versión vieja")

except ImportError:
    st.warning("Inicializando entorno de ejecución... esto solo pasará la primera vez.")
    instalar("streamlit>=1.40.0")
    instalar("openai")
    instalar("python-docx")
    instalar("fpdf")
    st.rerun() 

import openai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import io

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Actas EF - IES Lucía de Medrano", page_icon="📝")

# --- GESTIÓN DE ESTADO (SESSION STATE) ---
if 'grabaciones_guardadas' not in st.session_state:
    st.session_state.grabaciones_guardadas = []
if 'contador_micro' not in st.session_state:
    st.session_state.contador_micro = 0
if 'uploader_key' not in st.session_state:
    st.session_state.uploader_key = 0

# --- LÓGICA CORE (Blindada contra errores de codificación) ---

def transcribir_audio(audio_file, api_key):
    """
    Transcribe el audio asegurando que no existan errores de codificación ASCII.
    Técnica: Byte Stream Cloning.
    """
    client = openai.OpenAI(api_key=api_key)
    
    # 1. Rebobinar el archivo original por seguridad
    audio_file.seek(0)
    
    # 2. Detección segura de extensión (sin usar el nombre completo para evitar ASCII errors)
    # Por defecto usamos wav si no podemos determinarlo
    ext_segura = "wav"
    try:
        if hasattr(audio_file, 'name') and "." in audio_file.name:
            candidato = audio_file.name.split(".")[-1].lower()
            if candidato in ["mp3", "wav", "m4a", "mp4", "mpeg", "mpga", "webm"]:
                ext_segura = candidato
    except:
        pass # Ante cualquier error de lectura de nombre, fallamos a wav
    
    # 3. CREACIÓN DE NOMBRE SEGURO (ASCII PURO)
    # Ignoramos "Grabación.m4a" y usamos "input_file.m4a"
    nombre_ascii = f"input_file.{ext_segura}"
    
    # 4. CLONACIÓN EN MEMORIA (El paso crítico)
    # Leemos los bytes y los metemos en un buffer nuevo que no tiene metadatos del sistema operativo
    audio_bytes = audio_file.read()
    buffer_limpio = io.BytesIO(audio_bytes)
    buffer_limpio.name = nombre_ascii # Asignamos el nombre seguro al buffer
    
    # 5. Llamada a la API
    # Usamos temperature=0 para reducir alucinaciones en la transcripción fonética
    transcript = client.audio.transcriptions.create(
        model="whisper-1", 
        file=(nombre_ascii, buffer_limpio), 
        language="es",
        temperature=0 
    )
    return transcript.text

def generar_contenido_acta(transcripcion_completa, fecha, api_key):
    client = openai.OpenAI(api_key=api_key)
    
    # Prompt de Ingeniería para máxima fidelidad
    prompt_sistema = f"""
    Actúa como un secretario administrativo riguroso del Departamento de Educación Física.
    
    OBJETIVO: Convertir la transcripción en un ACTA OFICIAL FIDELIGNA.
    
    REGLAS DE SEGURIDAD (ANTI-ALUCINACIONES):
    1. VERACIDAD ABSOLUTA: Solo escribe lo que está en el texto. Si no se menciona un tema, no lo inventes.
    2. IDENTIDAD: Si no sabes quién habla, usa "Un docente" o "Un asistente". NUNCA inventes nombres propios (como Juan, María) si no se escuchan en el audio.
    3. ESTILO: 
       - Narración formal e impersonal para acuerdos generales.
       - CITA LITERAL OBLIGATORIA: Cuando alguien opine o diga "que conste en acta", escribe: D./Dña. [Nombre/Un docente] manifestó: "[Palabras textuales]".
    4. SILENCIOS/RUIDO: Si un audio no tiene contenido inteligible, ignóralo o indica "Tramo ininteligible".
    
    FORMATO DE SALIDA:
    - BLOQUE 1: "AUSENCIAS: [Lista o Ninguna]".
    - BLOQUE 2: Desarrollo de la sesión (Párrafos completos, no esquemas).
    """

    response = client.chat.completions.create(
        model="gpt-4o", 
        temperature=0.2, # Baja temperatura = Baja creatividad = Alta fidelidad
        messages=[
            {"role": "system", "content": prompt_sistema},
            {"role": "user", "content": f"Fecha: {fecha}. TRANSCRIPCIÓN BRUTA:\n\n{transcripcion_completa}"}
        ]
    )
    return response.choices[0].message.content

def crear_documento_word(contenido_ai, fecha):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Encabezado
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("Acta del Departamento de Educación Física del IES Lucía de Medrano")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph("") 

    # Fecha
    p_fecha = doc.add_paragraph()
    p_fecha.add_run("Fecha de la sesión: ").bold = True
    p_fecha.add_run(str(fecha))

    # Procesar texto
    lineas = contenido_ai.split('\n')
    texto_cuerpo = ""
    texto_asistentes = "Todos los componentes del Departamento de EF"
    
    for linea in lineas:
        if "AUSENCIAS:" in linea:
            if "Ninguna" not in linea:
                ausentes = linea.replace("AUSENCIAS:", "").strip()
                texto_asistentes += f", excepto {ausentes}"
        else:
            texto_cuerpo += linea + "\n"

    # Asistentes
    p_asist = doc.add_paragraph()
    p_asist.add_run("Asistentes: ").bold = True
    p_asist.add_run(texto_asistentes)
    doc.add_paragraph("") 
    
    # Cuerpo
    doc.add_heading('Desarrollo de la sesión:', level=2)
    doc.add_paragraph(texto_cuerpo.strip())
    doc.add_paragraph("") 

    # Cierre
    p_cierre = doc.add_paragraph()
    p_cierre.add_run("Y para que conste en acta y surta los efectos oportunos donde proceda firmo la siguiente.\n")
    p_cierre.add_run(f"En Salamanca a {fecha}")
    doc.add_paragraph("") 
    doc.add_paragraph("") 
    
    p_firma = doc.add_paragraph("EL JEFE DEL DEPARTAMENTO DE EDUCACIÓN FÍSICA")
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma.runs[0].bold = True

    bio = io.BytesIO()
    doc.save(bio)
    return bio

# --- INTERFAZ DE USUARIO ---
st.title("📝 Generador de Actas - Dpto. Educación Física")
st.markdown("**IES Lucía de Medrano** | *v3.0 Stable (ASCII Safe)*")

# 1. API KEY
with st.expander("🔐 Configuración", expanded=not st.session_state.get('api_ok', False)):
    api_key = st.text_input("Introduce tu API Key de OpenAI:", type="password")
    if api_key: st.session_state['api_ok'] = True

st.divider()
fecha_sesion = st.date_input("📅 Fecha de la sesión", date.today())

# 2. GESTOR DE CARGA
st.write("### 🎙️ Gestión de Audios")
st.caption("Sistema de procesamiento seguro: Los nombres de archivo se normalizan automáticamente.")

tab1, tab2 = st.tabs(["📂 1. Subir Archivos", "🎤 2. Grabar (Multi-toma)"])

with tab1:
    key_dinamica = f"uploader_{st.session_state.uploader_key}"
    archivos_subidos = st.file_uploader(
        "Arrastra archivos aquí", 
        type=["mp3", "m4a", "wav"], 
        accept_multiple_files=True,
        key=key_dinamica
    )

with tab2:
    st.write("Graba y pulsa 'Guardar'.")
    key_micro = f"micro_input_{st.session_state.contador_micro}"
    audio_temporal = st.audio_input("Microfono", key=key_micro)

    if audio_temporal is not None:
        st.success("✅ Audio capturado")
        col_a, col_b = st.columns(2)
        with col_a:
            st.audio(audio_temporal)
        with col_b:
            if st.button("💾 GUARDAR Y LIMPIAR", type="primary"):
                timestamp = date.today().strftime("%H-%M-%S")
                # Forzamos nombre seguro al guardar en memoria
                audio_temporal.name = f"mic_rec_{timestamp}.wav" 
                st.session_state.grabaciones_guardadas.append(audio_temporal)
                st.session_state.contador_micro += 1
                st.rerun()

# 3. LISTADO
st.divider()
st.subheader("🎧 Audios listos para procesar")

lista_total = []
if archivos_subidos:
    lista_total.extend(archivos_subidos)
lista_total.extend(st.session_state.grabaciones_guardadas)

count = len(lista_total)

if count == 0:
    st.markdown("*La lista está vacía.*")
else:
    for i, audio in enumerate(lista_total):
        # Mostramos nombre original en la UI pero internamente está controlado
        nombre_display = getattr(audio, 'name', f"Audio_{i+1}")
        st.text(f"{i+1}. {nombre_display}")

# 4. EJECUCIÓN
st.divider()
boton_finalizar = st.button(
    f"✅ PROCESAR {count} AUDIOS Y GENERAR ACTA", 
    type="primary", 
    use_container_width=True,
    disabled=(count == 0)
)

if boton_finalizar:
    if not api_key:
        st.error("⚠️ Falta la API Key.")
    elif count > 20:
        st.error(f"⚠️ Límite de 20 archivos superado.")
    else:
        transcripcion_total = ""
        barra = st.progress(0, text="Iniciando...")
        error_fatal = False
        
        try:
            # Fase 1: Transcripción Robusta
            for i, archivo in enumerate(lista_total):
                barra.progress((i / count) * 0.7, text=f"Transcribiendo audio {i+1}/{count}...")
                try:
                    texto = transcribir_audio(archivo, api_key)
                    transcripcion_total += f"\n--- Intervención {i+1} ---\n{texto}\n"
                except Exception as e:
                    # Capturamos error individual pero intentamos seguir
                    st.error(f"Error procesando audio {i+1}: {e}")
                    error_fatal = True
            
            # Verificación de contenido
            if not transcripcion_total.strip():
                st.error("Error crítico: No se obtuvo texto de ningún audio.")
            else:
                # Fase 2: Redacción GPT-4
                barra.progress(0.80, text="Generando redacción oficial sin invenciones...")
                contenido = generar_contenido_acta(transcripcion_total, fecha_sesion, api_key)
                
                # Fase 3: Word
                barra.progress(0.95, text="Creando archivo descargable...")
                doc = crear_documento_word(contenido, fecha_sesion)
                
                barra.progress(1.0, text="¡Proceso completado!")
                st.balloons()
                
                if error_fatal:
                    st.warning("El acta se generó parcialmente (algunos audios fallaron).")
                else:
                    st.success("🎉 Acta generada con éxito.")
                
                st.download_button(
                    label="📥 DESCARGAR ACTA (.DOCX)",
                    data=doc.getvalue(),
                    file_name=f"Acta_EF_{fecha_sesion}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )

        except Exception as e:
            st.error(f"Error del sistema: {e}")

# 5. BORRADO SEGURO
st.write("---")
st.write("### 🗑️ Zona de Limpieza")

if 'mostrar_confirmacion' not in st.session_state:
    st.session_state.mostrar_confirmacion = False

col_clean1, col_clean2 = st.columns([3, 1])
with col_clean1:
    st.caption("Una vez descargada el acta, elimina los archivos para proteger la privacidad.")

with col_clean2:
    if st.button("Borrar Archivos", type="secondary"):
        st.session_state.mostrar_confirmacion = True

if st.session_state.mostrar_confirmacion:
    st.warning("⚠️ ¿Confirmas el borrado de todos los audios?")
    col_conf1, col_conf2 = st.columns(2)
    with col_conf1:
        if st.button("❌ Cancelar"):
            st.session_state.mostrar_confirmacion = False
            st.rerun()
    with col_conf2:
        if st.button("✅ SÍ, BORRAR TODO"):
            st.session_state.grabaciones_guardadas = []
            st.session_state.contador_micro = 0
            st.session_state.uploader_key += 1
            st.session_state.mostrar_confirmacion = False
            st.rerun()
